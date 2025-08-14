from flask import Flask, render_template, request, send_file, flash, redirect, url_for
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime, timedelta
import os
from openpyxl import Workbook, load_workbook
import locale
from num2words import num2words

app = Flask(__name__)
app.secret_key = 'your_secret_key_here'
app.config['UPLOAD_FOLDER'] = 'static/generated_docs'
app.config['DATA_FOLDER'] = 'data'

# Set locale for French number formatting
try:
    locale.setlocale(locale.LC_ALL, 'fr_FR.UTF-8')
except:
    locale.setlocale(locale.LC_ALL, 'french')

app.jinja_env.globals.update(datetime=datetime, timedelta=timedelta)

# Ensure directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['DATA_FOLDER'], exist_ok=True)

def init_excel_files():
    clients_path = os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx')
    items_path = os.path.join(app.config['DATA_FOLDER'], 'items.xlsx')
    
    if not os.path.exists(clients_path):
        wb = Workbook()
        ws = wb.active
        ws.append(['Client ID', 'Name', 'Address', 'Email', 'Phone', 'Fiscal ID', 'AI Number', 'RC Number'])
        wb.save(clients_path)
    
    if not os.path.exists(items_path):
        wb = Workbook()
        ws = wb.active
        ws.append(['Item ID', 'Code', 'Description', 'Unit Price', 'Category'])
        wb.save(items_path)

init_excel_files()

def get_clients():
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx'))
        ws = wb.active
        clients = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                clients.append({
                    'id': row[0],
                    'name': row[1],
                    'address': row[2],
                    'email': row[3] if len(row) > 3 else '',
                    'phone': row[4] if len(row) > 4 else '',
                    'fiscal_id': row[5] if len(row) > 5 else '',
                    'ai_number': row[6] if len(row) > 6 else '',
                    'rc_number': row[7] if len(row) > 7 else ''
                })
        return clients
    except Exception as e:
        flash(f'Error loading clients: {str(e)}', 'error')
        return []

def get_items():
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'items.xlsx'))
        ws = wb.active
        items = []
        for row in ws.iter_rows(min_row=2, values_only=True):
            if row and row[0]:
                unit_price = 0.0
                if len(row) > 3 and row[3] is not None:
                    try:
                        if isinstance(row[3], (int, float)):
                            unit_price = float(row[3])
                        else:
                            unit_price = float(str(row[3]).replace(',', '.'))
                    except (ValueError, TypeError):
                        unit_price = 0.0
                
                items.append({
                    'id': row[0],
                    'code': str(row[1]) if len(row) > 1 and row[1] else f"ITEM{row[0]:03d}",
                    'description': str(row[2]) if len(row) > 2 and row[2] else '',
                    'unit_price': unit_price,
                    'category': str(row[4]) if len(row) > 4 and row[4] else ''
                })
        return items
    except Exception as e:
        flash(f'Error loading items: {str(e)}', 'error')
        return []

def add_client(data):
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx'))
        ws = wb.active
        new_id = ws.max_row
        ws.append([
            new_id,
            data['name'],
            data['address'],
            data.get('email', ''),
            data.get('phone', ''),
            data.get('fiscal_id', ''),
            data.get('ai_number', ''),
            data.get('rc_number', '')
        ])
        wb.save(os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx'))
        return True
    except Exception as e:
        flash(f'Error adding client: {str(e)}', 'error')
        return False

def add_item(data):
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'items.xlsx'))
        ws = wb.active
        new_id = ws.max_row
        
        unit_price = 0.0
        if 'unit_price' in data:
            try:
                unit_price = float(str(data['unit_price']).replace(',', '.'))
            except (ValueError, TypeError):
                flash('Prix unitaire invalide. Utilisez un nombre (ex: 1500.00)', 'error')
                return False
        
        code = data.get('code', '').strip()
        if not code:
            code = f"ITEM{new_id:03d}"
            
        ws.append([
            new_id,
            code,
            data['description'],
            unit_price,
            data.get('category', '').strip()
        ])
        wb.save(os.path.join(app.config['DATA_FOLDER'], 'items.xlsx'))
        return True
    except Exception as e:
        flash(f'Error adding item: {str(e)}', 'error')
        return False

def delete_client(client_id):
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx'))
        ws = wb.active
        for idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            if row[0].value == client_id:
                ws.delete_rows(idx)
                break
        wb.save(os.path.join(app.config['DATA_FOLDER'], 'clients.xlsx'))
        return True
    except Exception as e:
        flash(f'Error deleting client: {str(e)}', 'error')
        return False

def delete_item(item_id):
    try:
        wb = load_workbook(os.path.join(app.config['DATA_FOLDER'], 'items.xlsx'))
        ws = wb.active
        for idx, row in enumerate(ws.iter_rows(min_row=2), start=2):
            if row[0].value == item_id:
                ws.delete_rows(idx)
                break
        wb.save(os.path.join(app.config['DATA_FOLDER'], 'items.xlsx'))
        return True
    except Exception as e:
        flash(f'Error deleting item: {str(e)}', 'error')
        return False

def generate_document(filepath, doc_type, company_info, client_info, items, payment_info):
    doc = Document()
    
    # Set document margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
    
    # Add company header
    header = doc.add_paragraph()
    header_run = header.add_run("SARL PROTRONIC INTERNATIONAL\n")
    header_run.bold = True
    header_run.font.size = Pt(14)
    
    header.add_run("Distribution\n")
    header.add_run("03, Rue Louise de Betignie - Alger\n")
    header.add_run("Tél: 021.64.63.60 / 61  Fax: 021.64.66.65\n")
    
    # Add bank information
    bank_info = doc.add_paragraph()
    bank_info.add_run("Compte : Société Générale Algérie\n")
    bank_info.add_run("RIB : 0210002113000003042\n")
    bank_info.add_run("e-mail : direction@protronic-dz.com\n")
    bank_info.add_run("Site web : www.protronic-dz.com\n")
    
    # Add separator line
    doc.add_paragraph("_" * 80)
    
    # Add company legal information
    legal_info = doc.add_paragraph()
    legal_info.add_run("Rc : 998 0008022\n")
    legal_info.add_run("A1 : 16017486401\n")
    legal_info.add_run("Id Fiscal : 09991600803277\n")
    legal_info.add_run("NIS : 099616010516338\n")
    
    # Document title mapping
    doc_titles = {
        'proforma': 'PROFORMA',
        'invoice': 'FACTURE',
        'quote': 'DEVIS',
        'discounted_invoice': 'FACTURE AVEC REMISE',
        'tax_exempt': 'FACTURE HORS TAXE'
    }
    
    title = doc.add_paragraph()
    title_run = title.add_run(f"{doc_titles.get(doc_type, 'DOCUMENT')} {company_info['doc_number']}\n")
    title_run.bold = True
    title_run.font.size = Pt(14)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Add payment method and client info
    payment_method = doc.add_paragraph()
    payment_method.add_run(f"Mode de Paiement : {payment_info['method']}\n")
    payment_method.add_run(f"{client_info['name']}\n")
    payment_method.add_run(f"par : {client_info['address']}\n")
    payment_method.add_run(f"IF : {client_info['fiscal_id']} AI: {client_info['ai_number']} RC : {client_info['rc_number']}\n")
    
    # Add items table
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set column widths
    widths = [0.5, 1.5, 4, 1, 1, 1, 1]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
    
    # Add table headers
    hdr_cells = table.rows[0].cells
    headers = ['N°', 'CODE', 'DÉSIGNATION', 'QUANTITÉ', 'PU HT', 'RIS.%', 'MONTANT HT']
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        hdr_cells[i].paragraphs[0].runs[0].bold = True
    
    # Add items to table
    for idx, item in enumerate(items, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = item['code']
        row_cells[2].text = item['description']
        row_cells[3].text = str(item['quantity'])
        row_cells[4].text = locale.format_string('%.2f', item['unit_price'], grouping=True)
        row_cells[5].text = str(item['discount']) if item['discount'] else '-'
        row_cells[6].text = locale.format_string('%.2f', item['total'], grouping=True)
    
    # Add note about units
    doc.add_paragraph(f"NB. UV : {items[0]['quantity'] if len(items) == 1 else 'VARIÉ'}")
    
    # Calculate all financial values
    subtotal = sum(item['total'] for item in items)
    discount = payment_info.get('discount', 0)
    discount_amount = subtotal * (discount / 100)
    net_subtotal = subtotal - discount_amount
    tax_rate = 0 if doc_type == 'tax_exempt' else payment_info.get('tax_rate', 19)
    tax_amount = net_subtotal * (tax_rate / 100)
    total = net_subtotal + tax_amount
    
    # Format total in French words
    try:
        total_in_words = num2words(total, lang='fr') + " DINARS"
        total_in_words = total_in_words.upper()
    except:
        total_in_words = f"{total:.2f} DINARS"
    
    # Add total in words
    doc.add_paragraph(f"Arrêtée la présente {doc_titles.get(doc_type, 'document').lower()} à la somme de : {total_in_words}\n")
    
    # Add financial summary table
    totals_table = doc.add_table(rows=5, cols=2)
    totals_table.style = 'Table Grid'
    totals_table.columns[0].width = Inches(3)
    totals_table.columns[1].width = Inches(2)
    
    # Add financial rows
    rows = [
        ("TOTAL HT", subtotal),
        (f"REMISE {discount}%", discount_amount) if discount > 0 else None,
        ("NET HT", net_subtotal),
        (f"TVA {tax_rate}%", tax_amount) if doc_type != 'tax_exempt' else None,
        ("TIMBRE", 0),
        ("NET A PAYER", total)
    ]
    
    row_idx = 0
    for row in rows:
        if row:
            cells = totals_table.rows[row_idx].cells
            cells[0].text = row[0]
            cells[1].text = locale.format_string('%.2f', row[1], grouping=True)
            if row[0] == "NET A PAYER":
                cells[1].paragraphs[0].runs[0].bold = True
            row_idx += 1
    
    # Add footer
    doc.add_paragraph("\n")
    footer = doc.add_paragraph()
    footer_run = footer.add_run("SARL PROTRONIC Commercial\n")
    footer_run.bold = True
    footer_run.font.size = Pt(12)
    
    warranty = doc.add_paragraph()
    warranty.add_run("Notre matériel est garanti une année contre tout vice de fabrication et est conforme aux normes internationales en vigueur.")
    warranty.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Save document
    doc.save(filepath)

@app.route('/', methods=['GET', 'POST'])
def index():
    doc_types = [
        ('proforma', 'Proforma'),
        ('invoice', 'Facture'),
        ('quote', 'Devis'),
        ('discounted_invoice', 'Facture avec remise'),
        ('tax_exempt', 'Facture hors taxe')
    ]
    
    if request.method == 'POST':
        try:
            doc_type = request.form['doc_type']
            client_id = int(request.form['client_id'])
            item_ids = request.form.getlist('item_ids')
            quantities = request.form.getlist('quantities')
            
            clients = get_clients()
            items_data = get_items()
            
            selected_client = next((c for c in clients if c['id'] == client_id), None)
            if not selected_client:
                flash('Client non trouvé', 'error')
                return redirect(url_for('index'))
            
            items = []
            for i, item_id in enumerate(item_ids):
                item = next((it for it in items_data if it['id'] == int(item_id)), None)
                if item:
                    quantity = float(quantities[i])
                    discount = float(request.form.get(f'discount_{item_id}', 0))
                    unit_price = float(item['unit_price'])
                    total = quantity * unit_price * (1 - discount/100)
                    
                    items.append({
                        'code': item['code'],
                        'description': item['description'],
                        'quantity': quantity,
                        'unit_price': unit_price,
                        'discount': discount,
                        'total': total
                    })
            
            if not items:
                flash('Aucun article sélectionné', 'error')
                return redirect(url_for('index'))
            
            # Calculate all financial values
            subtotal = sum(item['total'] for item in items)
            global_discount = float(request.form.get('global_discount', 0))
            discount_amount = subtotal * (global_discount / 100)
            net_subtotal = subtotal - discount_amount
            tax_rate = 19 if doc_type != 'tax_exempt' else 0
            tax_amount = net_subtotal * (tax_rate / 100)
            total = net_subtotal + tax_amount
            
            company_info = {
                'doc_number': f"FP{datetime.now().strftime('%y/%m')}/{datetime.now().strftime('%H%M')}",
                'doc_date': datetime.now().strftime('%d/%m/%Y')
            }
            
            payment_info = {
                'method': request.form.get('payment_method', 'ESPÈCES'),
                'discount': global_discount,
                'tax_rate': tax_rate
            }
            
            filename = f"{doc_type.upper()}_{company_info['doc_number'].replace('/', '_')}.docx"
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            
            generate_document(
                filepath,
                doc_type,
                company_info,
                {
                    'name': selected_client['name'],
                    'address': selected_client['address'],
                    'fiscal_id': selected_client.get('fiscal_id', ''),
                    'ai_number': selected_client.get('ai_number', ''),
                    'rc_number': selected_client.get('rc_number', '')
                },
                items,
                payment_info
            )
            
            return render_template('success.html', 
                                filename=filename,
                                doc_type=doc_type,
                                items=items,
                                subtotal=subtotal,
                                discount_amount=discount_amount,
                                net_subtotal=net_subtotal,
                                tax_amount=tax_amount,
                                total=total,
                                tax_rate=tax_rate,
                                global_discount=global_discount)
        
        except Exception as e:
            flash(f'Erreur: {str(e)}', 'error')
            return redirect(url_for('index'))
    
    return render_template('index.html', 
                         clients=get_clients(), 
                         items=get_items(),
                         doc_types=doc_types)

@app.route('/clients', methods=['GET', 'POST'])
def manage_clients():
    if request.method == 'POST':
        if 'delete' in request.form:
            if delete_client(int(request.form['client_id'])):
                flash('Client supprimé avec succès', 'success')
        else:
            if not request.form['name'] or not request.form['address']:
                flash('Le nom et l\'adresse sont obligatoires', 'error')
            elif add_client({
                'name': request.form['name'],
                'address': request.form['address'],
                'email': request.form.get('email', ''),
                'phone': request.form.get('phone', ''),
                'fiscal_id': request.form.get('fiscal_id', ''),
                'ai_number': request.form.get('ai_number', ''),
                'rc_number': request.form.get('rc_number', '')
            }):
                flash('Client ajouté avec succès', 'success')
        return redirect(url_for('manage_clients'))
    
    return render_template('clients.html', clients=get_clients())

@app.route('/items', methods=['GET', 'POST'])
def manage_items():
    if request.method == 'POST':
        if 'delete' in request.form:
            if delete_item(int(request.form['item_id'])):
                flash('Article supprimé avec succès', 'success')
        else:
            if not request.form['description'] or not request.form['unit_price']:
                flash('La description et le prix unitaire sont obligatoires', 'error')
            elif add_item({
                'code': request.form.get('code', '').strip(),
                'description': request.form['description'],
                'unit_price': request.form['unit_price'],
                'category': request.form.get('category', '').strip()
            }):
                flash('Article ajouté avec succès', 'success')
        return redirect(url_for('manage_items'))
    
    return render_template('items.html', items=get_items())

@app.route('/download/<filename>')
def download(filename):
    try:
        filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if not os.path.exists(filepath):
            flash('Fichier non trouvé', 'error')
            return redirect(url_for('index'))
        return send_file(filepath, as_attachment=True)
    except Exception as e:
        flash(f'Erreur de téléchargement: {str(e)}', 'error')
        return redirect(url_for('index'))

if __name__ == '__main__':
    app.run(debug=True)